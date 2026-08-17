"""Generate LoyaltyLion Connector documentation as DOCX."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ── Title Page ───────────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('LoyaltyLion → BigQuery Connector')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 178)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Deployment & Operations Guide')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Version 2.0 — August 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ── Table of Contents ────────────────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Architecture',
    '3. BigQuery Tables',
    '4. VM Deployment',
    '5. Scheduling',
    '6. Authentication',
    '7. Manual Operations',
    '8. Monitoring & Logs',
    '9. Troubleshooting',
    '10. File Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Overview ─────────────────────────────────────────────────────────────
doc.add_heading('1. Overview', level=1)

doc.add_paragraph(
    'The LoyaltyLion → BigQuery connector syncs loyalty program data from the '
    'LoyaltyLion v2 API into Google BigQuery. It runs on a GCP VM with automated '
    'twice-daily sync schedules via crontab.'
)

doc.add_heading('1.1 Key Features', level=2)
features = [
    '5 BigQuery tables: ll_customers, ll_activities, ll_transactions, ll_orders, ll_customer_balances',
    'Incremental sync (since_id) for activities, transactions, orders',
    'Updated-at sync for customers (catches point balance changes)',
    'Daily balance snapshots for trend analysis',
    'Streaming buffer fallback for BigQuery DELETE operations',
    'Full sync truncates all tables first to prevent duplicates',
    'Service account key auth prioritized over VM default credentials',
    '4-attempt retry on 429/5xx errors',
    'Rate limiting (0.5s between requests)',
    'Auto-creates tables with partitioning and clustering',
    'Git-based deployment with symlink for easy updates',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# ── 2. Architecture ─────────────────────────────────────────────────────────
doc.add_heading('2. Architecture', level=1)

doc.add_paragraph(
    'The connector follows the same pattern as other MSREAD connectors '
    '(Business Central, Facebook Ads, Shopify Inventory).'
)

arch_table = doc.add_table(rows=7, cols=2)
arch_table.style = 'Light Grid Accent 1'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Component', 'Details']
for i, h in enumerate(headers):
    arch_table.rows[0].cells[i].text = h
arch_data = [
    ('VM Name', 'loyaltylion (asia-southeast1-a)'),
    ('VM IP', '34.124.194.116'),
    ('VM User', 'msrit_msread'),
    ('BigQuery Project', 'msr-msia-sales-analysis'),
    ('BigQuery Dataset', 'loyaltylion'),
    ('GitHub Repo', 'https://github.com/msritmsread-collab/loyaltylion'),
]
for i, (k, v) in enumerate(arch_data, 1):
    arch_table.rows[i].cells[0].text = k
    arch_table.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_heading('2.1 Directory Layout', level=2)
doc.add_paragraph(
    '/opt/msread-loyaltylion  →  /home/msrit_msread/loyaltylion  (symlink)\n'
    '~/loyaltylion/                              # Git clone of the repo\n'
    '  ├── loyaltylion_connect.py                # API client\n'
    '  ├── loyaltylion_to_bigquery.py            # Pipeline orchestrator\n'
    '  ├── loyaltylion_credentials.json          # API key + BQ config (not in git)\n'
    '  ├── bigquery_credentials.json             # Service account key (not in git)\n'
    '  ├── loyaltylion_sync_state.json           # Runtime state (not in git)\n'
    '  ├── venv/                                 # Python virtual environment\n'
    '  └── logs/                                 # Cron log output\n'
    '      ├── sync_8am.log\n'
    '      └── sync_2pm.log',
    style='No Spacing'
)

doc.add_paragraph(
    'Data flow: LoyaltyLion API → loyaltylion_connect.py (API client) → '
    'loyaltylion_to_bigquery.py (pipeline) → BigQuery tables'
)

# ── 3. BigQuery Tables ──────────────────────────────────────────────────────
doc.add_heading('3. BigQuery Tables', level=1)

tables_info = [
    ('ll_customers', 'updated_at', 'email', 'updated_at_min', 'Catches point balance changes via DELETE+re-insert'),
    ('ll_activities', 'created_at', 'customer_id, state', 'since_id', 'Immutable events — fetches only new records'),
    ('ll_transactions', 'created_at', 'customer_id, transaction_type', 'since_id', 'Point transactions (activity, claimed_reward, adjustment)'),
    ('ll_orders', 'created_at', 'customer_id, state', 'since_id', 'Order loyalty data'),
    ('ll_customer_balances', 'snapshot_date', 'merchant_id, snapshot_date', 'Daily snapshot', 'Daily point balance snapshot for trend tracking'),
]

tbl = doc.add_table(rows=len(tables_info)+1, cols=5)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Table', 'Partition', 'Cluster', 'Incremental', 'Notes']
for i, h in enumerate(hdrs):
    tbl.rows[0].cells[i].text = h
for i, row_data in enumerate(tables_info, 1):
    for j, val in enumerate(row_data):
        tbl.rows[i].cells[j].text = val

# ── 4. VM Deployment ────────────────────────────────────────────────────────
doc.add_heading('4. VM Deployment', level=1)

doc.add_heading('4.1 One-Time Setup', level=2)
doc.add_paragraph(
    '# 1. Clone the repo\n'
    'git clone https://github.com/msritmsread-collab/loyaltylion.git ~/loyaltylion\n\n'
    '# 2. Create symlink\n'
    'sudo ln -s /home/msrit_msread/loyaltylion /opt/msread-loyaltylion\n\n'
    '# 3. Copy credentials (not in git)\n'
    'cp /tmp/loyaltylion_credentials.json ~/loyaltylion/\n'
    'cp /tmp/bigquery_credentials.json ~/loyaltylion/\n'
    'cp /tmp/loyaltylion_sync_state.json ~/loyaltylion/\n\n'
    '# 4. Update credentials path\n'
    'sed -i \'s|/opt/msread-loyaltylion/bigquery_credentials.json|'
    '/home/msrit_msread/loyaltylion/bigquery_credentials.json|\' '
    '~/loyaltylion/loyaltylion_credentials.json\n\n'
    '# 5. Create venv and install dependencies\n'
    'python3 -m venv ~/loyaltylion/venv\n'
    '~/loyaltylion/venv/bin/pip install -r ~/loyaltylion/requirements.txt\n\n'
    '# 6. Create log directory\n'
    'mkdir -p ~/loyaltylion/logs',
    style='No Spacing'
)

doc.add_heading('4.2 Alternative: Deploy Script', level=2)
doc.add_paragraph(
    'For fresh VMs, use the deploy script (sets up venv, cron, systemd timers):\n\n'
    'cd ~/loyaltylion\n'
    'bash deploy_loyaltylion.sh',
    style='No Spacing'
)

doc.add_heading('4.3 Upload Credentials', level=2)
doc.add_paragraph(
    'From your local machine (PowerShell), upload the BigQuery service account key:\n\n'
    'scp "C:\\path\\to\\bigquery_credentials.json" msrit_msread@34.124.194.116:~/loyaltylion/\n\n'
    'The loyaltylion_credentials.json contains the API key and BigQuery config. '
    'The bigquery_credentials.json is the GCP service account key for BigQuery access.\n\n'
    '⚠️ Never commit these files to git. They are excluded in .gitignore.'
)

# ── 5. Scheduling ───────────────────────────────────────────────────────────
doc.add_heading('5. Scheduling', level=1)

doc.add_paragraph('Crontab runs twice daily (Malaysia time, UTC+8):')

sched_table = doc.add_table(rows=3, cols=4)
sched_table.style = 'Light Grid Accent 1'
sched_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Schedule', 'UTC Time', 'MYT Time', 'Type']):
    sched_table.rows[0].cells[i].text = h
sched_data = [
    ('0 0 * * *', '00:00', '8:00 AM', 'Incremental'),
    ('0 6 * * *', '06:00', '2:00 PM', 'Incremental'),
]
for i, row_data in enumerate(sched_data, 1):
    for j, val in enumerate(row_data):
        sched_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_heading('5.1 Setup Crontab', level=2)
doc.add_paragraph(
    '# Install cron if not present\n'
    'sudo apt-get update && sudo apt-get install -y cron\n'
    'sudo systemctl enable cron && sudo systemctl start cron\n\n'
    '# Set up crontab\n'
    'echo \'0 0 * * * /home/msrit_msread/loyaltylion/venv/bin/python3 '
    '/home/msrit_msread/loyaltylion/loyaltylion_to_bigquery.py incremental '
    '>> /home/msrit_msread/loyaltylion/logs/sync_8am.log 2>&1\n'
    '0 6 * * * /home/msrit_msread/loyaltylion/venv/bin/python3 '
    '/home/msrit_msread/loyaltylion/loyaltylion_to_bigquery.py incremental '
    '>> /home/msrit_msread/loyaltylion/logs/sync_2pm.log 2>&1\' | crontab -\n\n'
    '# Verify\n'
    'crontab -l',
    style='No Spacing'
)

doc.add_heading('5.2 Systemd Timers (Alternative)', level=2)
doc.add_paragraph(
    'The deploy script also creates systemd timers as an alternative to cron.\n'
    'To enable:\n\n'
    'sudo systemctl enable --now loyaltylion-incremental.timer\n'
    'sudo systemctl enable --now loyaltylion-full.timer\n\n'
    'If using systemd, remove cron entries: crontab -r'
)

# ── 6. Authentication ────────────────────────────────────────────────────────
doc.add_heading('6. Authentication', level=1)

doc.add_paragraph('The connector uses a priority-based auth system:')

auth_table = doc.add_table(rows=3, cols=3)
auth_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Component', 'Priority 1 (tried first)', 'Priority 2 (fallback)']):
    auth_table.rows[0].cells[i].text = h
auth_data = [
    ('LoyaltyLion API Key', 'GCP Secret Manager: connector-loyaltylion-api-key', 'loyaltylion_credentials.json → api_key'),
    ('BigQuery', 'Service account key file (bigquery_credentials.json)', 'VM default credentials'),
]
for i, row_data in enumerate(auth_data, 1):
    for j, val in enumerate(row_data):
        auth_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Important: BigQuery auth prioritizes the service account key file over VM default credentials. '
    'This fixes a 403 ACCESS_TOKEN_SCOPE_INSUFFICIENT error that occurs because the VM\'s default '
    'service account lacks BigQuery scopes. The SA key file (myai-996@msr-msia-sales-analysis.iam.gserviceaccount) '
    'has the necessary BigQuery permissions.'
)

# ── 7. Manual Operations ───────────────────────────────────────────────────
doc.add_heading('7. Manual Operations', level=1)

doc.add_heading('7.1 Manual Sync', level=2)
doc.add_paragraph(
    '# Incremental sync (recommended for daily use)\n'
    '~/loyaltylion/venv/bin/python3 ~/loyaltylion/loyaltylion_to_bigquery.py incremental\n\n'
    '# Full sync (truncates all tables first, re-fetches everything)\n'
    '~/loyaltylion/venv/bin/python3 ~/loyaltylion/loyaltylion_to_bigquery.py full',
    style='No Spacing'
)

doc.add_heading('7.2 Update Code', level=2)
doc.add_paragraph(
    'cd ~/loyaltylion && git pull\n\n'
    'No need to reinstall dependencies unless requirements.txt changed.\n'
    'If it did:\n'
    '~/loyaltylion/venv/bin/pip install -r ~/loyaltylion/requirements.txt',
    style='No Spacing'
)

doc.add_heading('7.3 Reset Sync State', level=2)
doc.add_paragraph(
    'To force a full re-sync, delete the sync state file:\n\n'
    'rm ~/loyaltylion/loyaltylion_sync_state.json\n\n'
    'Then run: ~/loyaltylion/venv/bin/python3 ~/loyaltylion/loyaltylion_to_bigquery.py full',
    style='No Spacing'
)

# ── 8. Monitoring & Logs ────────────────────────────────────────────────────
doc.add_heading('8. Monitoring & Logs', level=1)

doc.add_heading('8.1 Cron Logs', level=2)
doc.add_paragraph(
    '# 8:00 AM sync log\n'
    'tail -f ~/loyaltylion/logs/sync_8am.log\n\n'
    '# 2:00 PM sync log\n'
    'tail -f ~/loyaltylion/logs/sync_2pm.log',
    style='No Spacing'
)

doc.add_heading('8.2 Systemd Logs (if using timers)', level=2)
doc.add_paragraph(
    '# Follow incremental sync logs\n'
    'sudo journalctl -u loyaltylion-incremental -f\n\n'
    '# Follow full sync logs\n'
    'sudo journalctl -u loyaltylion-full -f\n\n'
    '# View last 50 lines\n'
    'sudo journalctl -u loyaltylion-incremental -n 50',
    style='No Spacing'
)

doc.add_heading('8.3 Check Timer Status', level=2)
doc.add_paragraph(
    '# Cron (primary)\n'
    'crontab -l\n\n'
    '# Systemd (if using timers)\n'
    'systemctl list-timers --no-pager | grep loyaltylion\n'
    'sudo systemctl status loyaltylion-incremental.timer',
    style='No Spacing'
)

doc.add_heading('8.4 Check BigQuery Row Counts', level=2)
doc.add_paragraph('Run in BigQuery Console:')
doc.add_paragraph(
    "SELECT table_id, row_count\n"
    "FROM `msr-msia-sales-analysis.loyaltylion.INFORMATION_SCHEMA.TABLES`\n"
    "ORDER BY table_id",
    style='No Spacing'
)

doc.add_heading('8.5 Check Sync State', level=2)
doc.add_paragraph(
    'cat ~/loyaltylion/loyaltylion_sync_state.json | python3 -m json.tool',
    style='No Spacing'
)

# ── 9. Troubleshooting ──────────────────────────────────────────────────────
doc.add_heading('9. Troubleshooting', level=1)

trouble_table = doc.add_table(rows=8, cols=2)
trouble_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Issue', 'Solution']):
    trouble_table.rows[0].cells[i].text = h
trouble_data = [
    ('403 ACCESS_TOKEN_SCOPE_INSUFFICIENT (BigQuery)',
     'The VM default credentials lack BigQuery scope. The connector now prioritizes the SA key file (bigquery_credentials.json) over VM default creds. Ensure this file exists at ~/loyaltylion/bigquery_credentials.json.'),
    ('403 ACCESS_TOKEN_SCOPE_INSUFFICIENT (Secret Manager)',
     'Expected on this VM — it was created without --scopes=cloud-platform. The connector falls back to local credential files automatically. No action needed.'),
    ('Streaming buffer error on DELETE',
     'BigQuery cannot DELETE rows still in the streaming buffer. The pipeline catches this and falls back to TRUNCATE TABLE for both delete_updated_since and delete_for_date.'),
    ('FileNotFoundError: bigquery_credentials.json',
     'Upload the SA key file to ~/loyaltylion/bigquery_credentials.json via SCP from your local machine, or create it using the heredoc method.'),
    ('FileNotFoundError: loyaltylion_credentials.json',
     'Copy from /tmp if it was backed up, or recreate from the example template.'),
    ('429 Rate Limited',
     'The connector retries up to 4 times with exponential backoff. The rate_limit_seconds=0.5 setting throttles requests.'),
    ('Duplicate rows after full sync',
     'Full sync now truncates all tables before re-fetching. If you see duplicates, ensure you are running the latest code (git pull).'),
]
for i, (issue, sol) in enumerate(trouble_data, 1):
    trouble_table.rows[i].cells[0].text = issue
    trouble_table.rows[i].cells[1].text = sol

# ── 10. File Reference ──────────────────────────────────────────────────────
doc.add_heading('10. File Reference', level=1)

files_table = doc.add_table(rows=9, cols=3)
files_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['File', 'Location (VM)', 'Purpose']):
    files_table.rows[0].cells[i].text = h
files_data = [
    ('loyaltylion_connect.py', '~/loyaltylion/', 'API client — Bearer auth, cursor pagination, rate limiting, flatten functions'),
    ('loyaltylion_to_bigquery.py', '~/loyaltylion/', 'Pipeline orchestrator — 5 tables, incremental/full sync, streaming buffer fallback, SA key auth priority'),
    ('loyaltylion_credentials.json', '~/loyaltylion/', 'API key + BigQuery config (⚠️ secrets, not in git)'),
    ('bigquery_credentials.json', '~/loyaltylion/', 'Service account key for BigQuery (⚠️ secrets, not in git)'),
    ('loyaltylion_sync_state.json', '~/loyaltylion/', 'Runtime state — since_id watermarks, last run timestamp (not in git)'),
    ('requirements.txt', '~/loyaltylion/', 'Python dependencies'),
    ('deploy_loyaltylion.sh', '~/loyaltylion/', 'VM deployment script — venv, cron, systemd timers'),
    ('loyaltylion_README.md', '~/loyaltylion/', 'Quick reference documentation'),
]
for i, row_data in enumerate(files_data, 1):
    for j, val in enumerate(row_data):
        files_table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.size = Pt(10)

# Save
output_path = r'C:\Users\MUHAMMADHAZWANBINOTH\OneDrive - msread.com.my\Documents\MYAI\LoyaltyLion_BigQuery_Connector_Documentation.docx'
doc.save(output_path)
print(f'Saved to {output_path}')